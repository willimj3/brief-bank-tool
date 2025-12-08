"""
Generate synthetic sample briefs for testing the Brief Bank tool.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_brief(filename: str, content: dict):
    """Create a DOCX brief from content dictionary."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Caption
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(content['court'].upper())
    run.bold = True

    doc.add_paragraph()

    # Case name
    case_para = doc.add_paragraph()
    case_para.add_run(content['case_name']).bold = True
    case_para.add_run(f"\n\nCase No. {content['case_number']}")
    case_para.add_run(f"\n\n{content['document_title'].upper()}")

    doc.add_paragraph("_" * 50)
    doc.add_paragraph()

    # Sections
    for section in content['sections']:
        # Section heading
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(section['title'])
        run.bold = True

        # Section content
        for para_text in section['paragraphs']:
            para = doc.add_paragraph(para_text)
            para.paragraph_format.line_spacing = 2.0
            para.paragraph_format.first_line_indent = Inches(0.5)

        doc.add_paragraph()

    # Save
    output_path = Path(__file__).parent / filename
    doc.save(output_path)
    print(f"Created: {output_path}")


# Sample Brief 1: Motion to Dismiss - 12(b)(6) - Contract Case
brief1 = {
    'court': 'United States District Court\nNorthern District of California',
    'case_name': 'TECHSTART INNOVATIONS, INC., Plaintiff,\nv.\nGLOBAL VENTURES LLC, Defendant.',
    'case_number': '3:24-cv-01234-WHO',
    'document_title': "Defendant's Motion to Dismiss Pursuant to Fed. R. Civ. P. 12(b)(6)",
    'sections': [
        {
            'title': 'I. INTRODUCTION',
            'paragraphs': [
                "Defendant Global Ventures LLC respectfully moves this Court to dismiss Plaintiff TechStart Innovations, Inc.'s Complaint pursuant to Federal Rule of Civil Procedure 12(b)(6) for failure to state a claim upon which relief can be granted. Plaintiff's breach of contract claim fails because the alleged agreement lacks the essential terms necessary to form an enforceable contract under California law, and Plaintiff's fraud claim fails to meet the heightened pleading standard of Rule 9(b).",
                "For the reasons set forth below, Defendant respectfully requests that this Court dismiss the Complaint in its entirety with prejudice."
            ]
        },
        {
            'title': 'II. STATEMENT OF FACTS',
            'paragraphs': [
                "According to the Complaint, Plaintiff TechStart Innovations, Inc. is a Delaware corporation with its principal place of business in San Francisco, California. Compl. ¶ 1. Defendant Global Ventures LLC is a Delaware limited liability company with its principal place of business in Palo Alto, California. Compl. ¶ 2.",
                "Plaintiff alleges that in March 2023, the parties engaged in discussions regarding a potential investment by Defendant in Plaintiff's Series B financing round. Compl. ¶ 10. Plaintiff further alleges that during these discussions, Defendant's representatives made oral statements expressing interest in investing 'approximately $5 million' in Plaintiff's company. Compl. ¶ 12.",
                "Plaintiff contends that based on these discussions, it declined to pursue other potential investors and proceeded with business expansion plans. Compl. ¶ 15. When Defendant ultimately did not make the anticipated investment, Plaintiff filed this action alleging breach of contract and fraud. Compl. ¶¶ 20-35."
            ]
        },
        {
            'title': 'III. LEGAL STANDARD',
            'paragraphs': [
                "A motion to dismiss under Rule 12(b)(6) tests the legal sufficiency of the claims alleged in the complaint. Navarro v. Block, 250 F.3d 729, 732 (9th Cir. 2001). To survive a motion to dismiss, a complaint must contain 'enough facts to state a claim to relief that is plausible on its face.' Bell Atl. Corp. v. Twombly, 550 U.S. 544, 570 (2007). A claim is facially plausible 'when the plaintiff pleads factual content that allows the court to draw the reasonable inference that the defendant is liable for the misconduct alleged.' Ashcroft v. Iqbal, 556 U.S. 662, 678 (2009).",
                "While the Court must accept all factual allegations in the complaint as true and construe them in the light most favorable to the plaintiff, it need not accept 'legal conclusions' or 'threadbare recitals of the elements of a cause of action, supported by mere conclusory statements.' Id. at 678. Courts are 'not bound to accept as true a legal conclusion couched as a factual allegation.' Twombly, 550 U.S. at 555.",
                "Claims sounding in fraud must also satisfy the heightened pleading requirements of Federal Rule of Civil Procedure 9(b), which requires that 'the circumstances constituting fraud or mistake shall be stated with particularity.' Fed. R. Civ. P. 9(b). This means the plaintiff must state 'the who, what, when, where, and how' of the alleged fraud. Vess v. Ciba-Geigy Corp. USA, 317 F.3d 1097, 1106 (9th Cir. 2003)."
            ]
        },
        {
            'title': 'IV. ARGUMENT',
            'paragraphs': [
                "A. Plaintiff's Breach of Contract Claim Fails Because No Enforceable Contract Exists",
                "To state a claim for breach of contract under California law, a plaintiff must allege: (1) the existence of a contract; (2) plaintiff's performance or excuse for nonperformance; (3) defendant's breach; and (4) resulting damages. Oasis West Realty, LLC v. Goldman, 51 Cal. 4th 811, 821 (2011). Plaintiff's claim founders on the first element—the existence of an enforceable contract.",
                "Under California law, 'an enforceable contract requires mutual consent, which cannot exist unless the parties agree upon the same thing in the same sense.' Bustamante v. Intuit, Inc., 141 Cal. App. 4th 199, 213 (2006). Moreover, 'a contract must be sufficiently definite for the court to ascertain the parties' obligations and to determine whether those obligations have been performed or breached.' Ersa Grae Corp. v. Fluor Corp., 1 Cal. App. 4th 613, 623 (1991).",
                "Here, Plaintiff's own allegations demonstrate that no enforceable contract was formed. The Complaint alleges only that Defendant expressed 'interest' in investing 'approximately $5 million.' Compl. ¶ 12. This vague expression of interest lacks the essential terms necessary to form a binding investment agreement, including: the precise investment amount, the equity percentage or valuation, the closing conditions, representations and warranties, and numerous other material terms that are standard in venture financing transactions.",
                "California courts have consistently held that preliminary negotiations and expressions of interest do not constitute enforceable contracts. See Kruse v. Bank of America, 202 Cal. App. 3d 38, 59 (1988) ('An agreement to agree in the future is not a binding contract.'). The alleged oral statements here are precisely such preliminary expressions that cannot, as a matter of law, form the basis of a breach of contract claim.",
                "B. Plaintiff's Fraud Claim Fails to Meet Rule 9(b)'s Particularity Requirements",
                "Plaintiff's fraud claim is equally deficient. Federal Rule of Civil Procedure 9(b) requires fraud allegations to be stated with particularity, including 'an account of the time, place, and specific content of the false representations as well as the identities of the parties to the misrepresentations.' Swartz v. KPMG LLP, 476 F.3d 756, 764 (9th Cir. 2007).",
                "The Complaint fails to identify: (1) which specific representative of Defendant allegedly made the fraudulent statements; (2) the exact words used; (3) when these statements were made beyond the vague 'March 2023'; or (4) where the statements were made. Without these basic details, Defendant cannot adequately respond to the fraud allegations, which is precisely why Rule 9(b) exists.",
                "Furthermore, under California law, a fraud claim requires the plaintiff to allege justifiable reliance. OCM Principal Opportunities Fund, L.P. v. CIBC World Markets Corp., 157 Cal. App. 4th 835, 864 (2007). Sophisticated business parties like Plaintiff cannot justifiably rely on oral statements regarding a multi-million dollar investment transaction that would ordinarily be documented in detailed written agreements. See Hinesley v. Oakshade Town Center, 135 Cal. App. 4th 289, 301 (2005)."
            ]
        },
        {
            'title': 'V. CONCLUSION',
            'paragraphs': [
                "For the foregoing reasons, Defendant Global Ventures LLC respectfully requests that this Court grant its Motion to Dismiss and dismiss Plaintiff's Complaint in its entirety with prejudice.",
                "DATED: November 15, 2024",
                "Respectfully submitted,\n\nWILSON & MARTINEZ LLP\n\n/s/ Sarah J. Wilson\nSarah J. Wilson (SBN 198765)\nCounsel for Defendant Global Ventures LLC"
            ]
        }
    ]
}


# Sample Brief 2: Opposition to Motion for Summary Judgment - Employment Discrimination
brief2 = {
    'court': 'United States District Court\nNorthern District of California',
    'case_name': 'MARIA SANTOS, Plaintiff,\nv.\nACME TECHNOLOGY CORPORATION, Defendant.',
    'case_number': '3:23-cv-05678-JST',
    'document_title': "Plaintiff's Opposition to Defendant's Motion for Summary Judgment",
    'sections': [
        {
            'title': 'I. INTRODUCTION',
            'paragraphs': [
                "Plaintiff Maria Santos respectfully submits this Opposition to Defendant Acme Technology Corporation's Motion for Summary Judgment. The evidence in this case creates genuine disputes of material fact that must be resolved by a jury, not on summary judgment.",
                "Plaintiff, a 52-year-old Latina woman, was terminated from her position as Senior Software Engineer after 12 years of exemplary performance, just months after complaining to Human Resources about discriminatory comments made by her supervisor. Defendant claims the termination was part of a 'reduction in force,' but the evidence shows that Plaintiff was the only senior engineer terminated while younger, less experienced engineers were retained. These facts present classic circumstantial evidence of both age and national origin discrimination under Title VII and the ADEA."
            ]
        },
        {
            'title': 'II. STATEMENT OF FACTS',
            'paragraphs': [
                "Maria Santos began working at Acme Technology Corporation in January 2011 as a Software Engineer. Santos Decl. ¶ 3. Over her 12-year tenure, she received consistently positive performance reviews, with ratings of 'Exceeds Expectations' or 'Outstanding' every year from 2011 through 2022. Santos Decl. ¶ 5; Ex. A (Performance Reviews).",
                "In 2019, Plaintiff was promoted to Senior Software Engineer, a position she held until her termination. Santos Decl. ¶ 7. Her 2022 performance review, completed just three months before her termination, stated that she was 'a valuable team member whose technical expertise and mentorship have been invaluable to the department.' Ex. A at 45.",
                "In September 2022, Defendant hired Michael Thompson as Plaintiff's new direct supervisor. Thompson Depo. at 34:5-10. Within weeks of his arrival, Thompson began making comments about Plaintiff's age and national origin. On multiple occasions, he referred to Plaintiff as 'the old guard' and questioned whether she could 'keep up with the younger engineers.' Santos Decl. ¶ 12; Martinez Decl. ¶ 8 (co-worker corroborating statements).",
                "Thompson also made comments about Plaintiff's accent, stating in a team meeting that 'maybe if people spoke clearer English, we wouldn't have so many bugs in our code.' Santos Decl. ¶ 14; Martinez Decl. ¶ 10. Plaintiff reported these comments to Human Resources in October 2022. Santos Decl. ¶ 16; Ex. B (HR Complaint).",
                "On January 15, 2023, Plaintiff was informed that her position was being eliminated as part of a 'reduction in force.' Santos Decl. ¶ 20. However, the evidence shows that Plaintiff was the only senior engineer terminated. Three junior engineers hired within the previous two years—all under age 35—were retained. Ex. C (Termination Records); Ex. D (Employee Roster). Moreover, Defendant hired two new software engineers within 60 days of Plaintiff's termination. Ex. E (Hiring Records)."
            ]
        },
        {
            'title': 'III. LEGAL STANDARD',
            'paragraphs': [
                "Summary judgment is appropriate only where 'there is no genuine dispute as to any material fact and the movant is entitled to judgment as a matter of law.' Fed. R. Civ. P. 56(a). The Court must view the evidence in the light most favorable to the nonmoving party and draw all reasonable inferences in her favor. Scott v. Harris, 550 U.S. 372, 378 (2007).",
                "In employment discrimination cases, 'courts should exercise particular caution before granting summary judgment, because the ultimate question is one that can only be resolved through a searching inquiry—one that is most appropriately conducted by a factfinder, upon a full record.' Schnidrig v. Columbia Mach., Inc., 80 F.3d 1406, 1410 (9th Cir. 1996).",
                "Under the burden-shifting framework of McDonnell Douglas Corp. v. Green, 411 U.S. 792 (1973), once a plaintiff establishes a prima facie case of discrimination, the burden shifts to the defendant to articulate a legitimate, non-discriminatory reason for the adverse action. If the defendant does so, the burden shifts back to the plaintiff to show that the stated reason is pretextual. Villiarimo v. Aloha Island Air, Inc., 281 F.3d 1054, 1062 (9th Cir. 2002)."
            ]
        },
        {
            'title': 'IV. ARGUMENT',
            'paragraphs': [
                "A. Plaintiff Has Established a Prima Facie Case of Discrimination",
                "To establish a prima facie case of discrimination, Plaintiff must show: (1) she belongs to a protected class; (2) she was performing her job satisfactorily; (3) she suffered an adverse employment action; and (4) similarly situated employees outside her protected class were treated more favorably. Hawn v. Exec. Jet Mgmt., Inc., 615 F.3d 1151, 1156 (9th Cir. 2010).",
                "Each element is satisfied here. Plaintiff, a 52-year-old Latina woman, belongs to protected classes under both Title VII (national origin) and the ADEA (age over 40). Her consistent 'Exceeds Expectations' and 'Outstanding' performance ratings establish satisfactory job performance. Her termination is clearly an adverse employment action. And the retention of younger, less experienced engineers while Plaintiff was terminated satisfies the fourth element.",
                "B. Defendant's Stated Reason Is Pretextual",
                "Defendant claims Plaintiff was terminated as part of a legitimate reduction in force. However, pretext can be shown through evidence that the employer's explanation is 'unworthy of credence' or that 'a discriminatory reason more likely motivated the employer.' Texas Dep't of Cmty. Affairs v. Burdine, 450 U.S. 248, 256 (1981).",
                "Multiple facts support a finding of pretext here. First, the temporal proximity between Plaintiff's HR complaint in October 2022 and her termination in January 2023 supports an inference of retaliation. See Yartzoff v. Thomas, 809 F.2d 1371, 1376 (9th Cir. 1987) (proximity in time can support inference of causation).",
                "Second, Defendant's claim of a 'reduction in force' is undermined by the fact that it hired two new software engineers within 60 days of Plaintiff's termination. See Nidds v. Schindler Elevator Corp., 113 F.3d 912, 918 (9th Cir. 1997) (hiring after RIF supports inference of pretext).",
                "Third, Thompson's documented discriminatory comments—calling Plaintiff 'the old guard,' questioning whether she could 'keep up with younger engineers,' and making derogatory comments about her accent—constitute direct evidence of discriminatory animus. See Godwin v. Hunt Wesson, Inc., 150 F.3d 1217, 1221 (9th Cir. 1998) (discriminatory statements by decisionmakers are probative of discriminatory intent).",
                "Fourth, the fact that Plaintiff was the only senior engineer terminated while less experienced, younger engineers were retained raises a strong inference of discrimination. See Reeves v. Sanderson Plumbing Prods., Inc., 530 U.S. 133, 147 (2000) (comparative evidence is relevant to showing pretext)."
            ]
        },
        {
            'title': 'V. CONCLUSION',
            'paragraphs': [
                "The evidence creates genuine disputes of material fact regarding whether Defendant's stated reason for Plaintiff's termination was pretextual and whether discriminatory animus motivated the decision. These disputes must be resolved by a jury.",
                "For the foregoing reasons, Plaintiff respectfully requests that this Court deny Defendant's Motion for Summary Judgment.",
                "DATED: March 1, 2024",
                "Respectfully submitted,\n\nJOHNSON & PARK LLP\n\n/s/ David Park\nDavid Park (SBN 234567)\nCounsel for Plaintiff Maria Santos"
            ]
        }
    ]
}


# Sample Brief 3: Motion for Preliminary Injunction - Trade Secret
brief3 = {
    'court': 'United States District Court\nNorthern District of California',
    'case_name': 'NEXGEN PHARMACEUTICALS, INC., Plaintiff,\nv.\nDR. JAMES CHEN and BIOTECH SOLUTIONS LLC, Defendants.',
    'case_number': '3:24-cv-02468-EMC',
    'document_title': "Plaintiff's Motion for Preliminary Injunction",
    'sections': [
        {
            'title': 'I. INTRODUCTION',
            'paragraphs': [
                "Plaintiff NexGen Pharmaceuticals, Inc. moves for a preliminary injunction to prevent Defendants Dr. James Chen and BioTech Solutions LLC from using, disclosing, or otherwise exploiting NexGen's trade secrets and confidential information. Dr. Chen, NexGen's former Director of Research, left NexGen to join competitor BioTech Solutions, taking with him proprietary research data, formulations, and clinical trial results that NexGen spent over $50 million and five years developing.",
                "Without immediate injunctive relief, NexGen will suffer irreparable harm as Defendants exploit NexGen's trade secrets to bring a competing drug to market, destroying NexGen's competitive advantage and causing harm that cannot be remedied by monetary damages alone."
            ]
        },
        {
            'title': 'II. STATEMENT OF FACTS',
            'paragraphs': [
                "NexGen Pharmaceuticals is a biopharmaceutical company specializing in novel cancer treatments. Declaration of CEO Robert Martinez (\"Martinez Decl.\") ¶ 3. Since 2019, NexGen has invested over $50 million in developing NGX-401, a groundbreaking immunotherapy treatment for non-small cell lung cancer. Martinez Decl. ¶ 5.",
                "Dr. James Chen joined NexGen in 2020 as Director of Research, overseeing the NGX-401 development program. Martinez Decl. ¶ 8. As a condition of his employment, Dr. Chen signed a Confidentiality and Non-Disclosure Agreement (\"NDA\") prohibiting disclosure of NexGen's confidential information and trade secrets. Ex. A (NDA). He also signed an agreement assigning all inventions to NexGen. Ex. B (Invention Assignment Agreement).",
                "During his tenure, Dr. Chen had access to NexGen's most sensitive trade secrets, including: (1) the proprietary formulation of NGX-401; (2) clinical trial protocols and results; (3) manufacturing processes; and (4) regulatory submission strategies. Martinez Decl. ¶ 12.",
                "On August 1, 2024, Dr. Chen abruptly resigned from NexGen. Martinez Decl. ¶ 15. Two days later, BioTech Solutions announced that Dr. Chen had joined as their Chief Scientific Officer and that the company was 'accelerating development of a novel immunotherapy for lung cancer.' Ex. C (BioTech Press Release).",
                "Forensic analysis of Dr. Chen's NexGen-issued laptop revealed that in the two weeks before his resignation, he copied over 10,000 files to an external hard drive, including the complete NGX-401 formulation data, clinical trial results, and regulatory submission drafts. Declaration of IT Director Sarah Kim (\"Kim Decl.\") ¶¶ 8-12; Ex. D (Forensic Report)."
            ]
        },
        {
            'title': 'III. LEGAL STANDARD',
            'paragraphs': [
                "To obtain a preliminary injunction, the moving party must establish: (1) likelihood of success on the merits; (2) likelihood of irreparable harm in the absence of preliminary relief; (3) that the balance of equities tips in the movant's favor; and (4) that an injunction is in the public interest. Winter v. Natural Res. Def. Council, Inc., 555 U.S. 7, 20 (2008).",
                "In the Ninth Circuit, courts apply a 'sliding scale' approach, under which 'serious questions going to the merits and a balance of hardships that tips sharply towards the plaintiff can support issuance of a preliminary injunction, so long as the plaintiff also shows that there is a likelihood of irreparable injury and that the injunction is in the public interest.' Alliance for the Wild Rockies v. Cottrell, 632 F.3d 1127, 1135 (9th Cir. 2011).",
                "Preliminary injunctions are 'an extraordinary remedy that may only be awarded upon a clear showing that the plaintiff is entitled to such relief.' Winter, 555 U.S. at 22."
            ]
        },
        {
            'title': 'IV. ARGUMENT',
            'paragraphs': [
                "A. NexGen Is Likely to Succeed on the Merits of Its Trade Secret Claims",
                "Under the Defend Trade Secrets Act (\"DTSA\"), 18 U.S.C. § 1836, a plaintiff must show: (1) the existence of a trade secret; (2) misappropriation of the trade secret; and (3) the trade secret was used in interstate commerce. InteliClear, LLC v. ETC Global Holdings, Inc., 978 F.3d 653, 657-58 (9th Cir. 2020).",
                "NexGen's NGX-401 formulation, clinical data, and manufacturing processes constitute trade secrets under the DTSA. Trade secret protection extends to information that: (1) derives independent economic value from not being generally known; and (2) is subject to reasonable efforts to maintain its secrecy. 18 U.S.C. § 1839(3). NexGen's $50 million investment and five years of research demonstrate the substantial independent value of this information. Martinez Decl. ¶ 20. NexGen maintained secrecy through NDAs, access controls, and security protocols. Kim Decl. ¶¶ 15-18.",
                "The evidence of misappropriation is overwhelming. Dr. Chen copied 10,000 confidential files to an external drive immediately before joining a direct competitor—the very definition of misappropriation under 18 U.S.C. § 1839(5). See Waymo LLC v. Uber Techs., Inc., 870 F.3d 1350, 1361 (Fed. Cir. 2017) (downloading files before departure supports misappropriation claim).",
                "B. NexGen Will Suffer Irreparable Harm Absent an Injunction",
                "Trade secret misappropriation 'typically results in irreparable harm' because 'trade secrets, once exposed, are gone forever.' Faiveley Transp. Malmo AB v. Wabtec Corp., 559 F.3d 110, 118 (2d Cir. 2009); see also Apple Inc. v. Samsung Elecs. Co., 735 F.3d 1352, 1360 (Fed. Cir. 2013) (loss of competitive advantage constitutes irreparable harm).",
                "If Defendants are permitted to use NexGen's trade secrets, NexGen will lose its first-mover advantage in the lung cancer immunotherapy market, a loss that cannot be quantified or remedied through damages. Martinez Decl. ¶ 25. Moreover, once the trade secrets are disclosed or used, they lose their protected status permanently. Id. ¶ 26.",
                "C. The Balance of Equities Favors NexGen",
                "The balance of equities strongly favors NexGen. NexGen seeks only to protect information it developed through years of research and tens of millions of dollars in investment. Defendants, by contrast, seek to profit from information they acquired through breach of contract and breach of duty.",
                "An injunction will not prevent Defendants from competing—it will only prevent them from competing unfairly using NexGen's stolen trade secrets. See FTC v. Actavis, Inc., 570 U.S. 136, 153 (2013) (legitimate competition is encouraged; unfair competition is not).",
                "D. An Injunction Serves the Public Interest",
                "The public interest favors protecting trade secrets and enforcing confidentiality agreements. Ruckelshaus v. Monsanto Co., 467 U.S. 986, 1012 (1984) (trade secret protection serves important public interests). Permitting the misappropriation of trade secrets undermines incentives for innovation and investment in research and development."
            ]
        },
        {
            'title': 'V. CONCLUSION',
            'paragraphs': [
                "For the foregoing reasons, NexGen respectfully requests that this Court enter a preliminary injunction: (1) prohibiting Defendants from using, disclosing, or otherwise exploiting any NexGen trade secrets or confidential information; (2) requiring Defendants to return all NexGen materials in their possession; and (3) prohibiting Dr. Chen from working on any lung cancer immunotherapy projects at BioTech Solutions for the duration of this litigation.",
                "DATED: September 15, 2024",
                "Respectfully submitted,\n\nMORGAN LEWIS LLP\n\n/s/ Jennifer Adams\nJennifer Adams (SBN 345678)\nCounsel for Plaintiff NexGen Pharmaceuticals, Inc."
            ]
        }
    ]
}


# Sample Brief 4: Motion to Dismiss - Personal Jurisdiction
brief4 = {
    'court': 'United States District Court\nNorthern District of California',
    'case_name': 'PACIFIC COAST DISTRIBUTORS, INC., Plaintiff,\nv.\nMIDWEST MANUFACTURING CO., Defendant.',
    'case_number': '3:24-cv-03691-LB',
    'document_title': "Defendant's Motion to Dismiss for Lack of Personal Jurisdiction",
    'sections': [
        {
            'title': 'I. INTRODUCTION',
            'paragraphs': [
                "Defendant Midwest Manufacturing Co. moves to dismiss this action pursuant to Federal Rule of Civil Procedure 12(b)(2) for lack of personal jurisdiction. Midwest is an Ohio corporation with no offices, employees, property, or bank accounts in California. It does not advertise in California, has no registered agent in California, and has never been licensed to do business in California.",
                "Plaintiff's sole basis for asserting jurisdiction is that Midwest sold products to Plaintiff, a California company, through a single transaction initiated by Plaintiff. This is insufficient to establish either general or specific personal jurisdiction over Midwest in California."
            ]
        },
        {
            'title': 'II. STATEMENT OF FACTS',
            'paragraphs': [
                "Midwest Manufacturing Co. is an Ohio corporation with its principal place of business in Columbus, Ohio. Declaration of CEO Thomas Miller (\"Miller Decl.\") ¶ 2. Midwest manufactures industrial equipment and sells its products primarily to customers in the Midwest region of the United States. Miller Decl. ¶ 3.",
                "Midwest has no offices, facilities, employees, or real property in California. Miller Decl. ¶ 5. It has no bank accounts in California and is not registered to do business in California. Miller Decl. ¶ 6. Midwest does not advertise in California and has no sales representatives in California. Miller Decl. ¶ 7.",
                "In January 2024, Plaintiff Pacific Coast Distributors contacted Midwest to inquire about purchasing industrial valves for resale. Miller Decl. ¶ 10. After negotiations conducted entirely by email and telephone, Midwest agreed to sell Plaintiff 500 industrial valves for $150,000. Miller Decl. ¶ 11. The contract was negotiated and executed in Ohio, and the valves were shipped F.O.B. Columbus, Ohio. Miller Decl. ¶ 12; Ex. A (Purchase Order).",
                "Plaintiff now alleges that the valves were defective and brings claims for breach of contract, breach of warranty, and negligence. Compl. ¶¶ 15-30."
            ]
        },
        {
            'title': 'III. LEGAL STANDARD',
            'paragraphs': [
                "Federal courts may exercise personal jurisdiction over a defendant only if the defendant has 'certain minimum contacts with [the forum] such that the maintenance of the suit does not offend traditional notions of fair play and substantial justice.' Int'l Shoe Co. v. Washington, 326 U.S. 310, 316 (1945) (internal quotation marks omitted).",
                "There are two types of personal jurisdiction: general and specific. Daimler AG v. Bauman, 571 U.S. 117, 127 (2014). General jurisdiction exists only when the defendant's contacts with the forum are 'so continuous and systematic as to render them essentially at home in the forum State.' Id. at 139. Specific jurisdiction exists when the plaintiff's claims 'arise out of or relate to the defendant's contacts with the forum.' Bristol-Myers Squibb Co. v. Superior Ct., 582 U.S. 255, 262 (2017).",
                "When a defendant moves to dismiss for lack of personal jurisdiction, the plaintiff bears the burden of establishing that jurisdiction exists. Mavrix Photo, Inc. v. Brand Techs., Inc., 647 F.3d 1218, 1223 (9th Cir. 2011)."
            ]
        },
        {
            'title': 'IV. ARGUMENT',
            'paragraphs': [
                "A. This Court Lacks General Jurisdiction Over Midwest",
                "General jurisdiction requires that a defendant's contacts with the forum be 'so continuous and systematic as to render it essentially at home' in the forum state. Daimler, 571 U.S. at 139. For a corporation, the 'paradigm' forums for general jurisdiction are its place of incorporation and principal place of business. Id. at 137.",
                "Midwest is incorporated in Ohio and has its principal place of business in Ohio. Miller Decl. ¶ 2. It has no continuous or systematic contacts with California—no offices, no employees, no property, no bank accounts, and no ongoing business activities in the state. Miller Decl. ¶¶ 5-7. Under Daimler, general jurisdiction clearly does not exist.",
                "B. This Court Lacks Specific Jurisdiction Over Midwest",
                "The Ninth Circuit applies a three-prong test for specific jurisdiction: (1) the defendant must purposefully direct activities at the forum or purposefully avail itself of the forum's benefits; (2) the claim must arise out of or relate to the defendant's forum-related activities; and (3) the exercise of jurisdiction must be reasonable. Axiom Foods, Inc. v. Acerchem Int'l, Inc., 874 F.3d 1064, 1068 (9th Cir. 2017).",
                "Midwest did not purposefully direct any activities at California. The undisputed evidence shows that Plaintiff initiated the transaction by contacting Midwest in Ohio. Miller Decl. ¶ 10. Midwest did not solicit Plaintiff's business, advertise in California, or reach out to California customers. The contract was negotiated and executed in Ohio, and the goods were shipped F.O.B. Ohio. Miller Decl. ¶¶ 11-12.",
                "The Supreme Court has repeatedly held that 'a defendant's relationship with a plaintiff or third party, standing alone, is an insufficient basis for jurisdiction.' Walden v. Fiore, 571 U.S. 277, 286 (2014). The mere fact that a plaintiff is a California company does not establish personal jurisdiction over an out-of-state defendant. Id.",
                "Similarly, the Ninth Circuit has held that 'entering into a contract with a forum resident, without more, does not establish the requisite minimum contacts.' Roth v. Garcia Marquez, 942 F.2d 617, 621 (9th Cir. 1991). Here, Midwest's only contact with California is a single sale to a California buyer—a 'random, fortuitous, or attenuated' contact insufficient to support jurisdiction. Burger King Corp. v. Rudzewicz, 471 U.S. 462, 475 (1985).",
                "The stream of commerce doctrine does not help Plaintiff. Under Asahi Metal Industry Co. v. Superior Court, 480 U.S. 102 (1987), placing products into the stream of commerce with awareness that they might reach the forum state is insufficient to establish jurisdiction. Here, Midwest shipped products directly to Plaintiff; it did not place them into any stream of commerce. Moreover, this was a single, isolated transaction, not an ongoing business relationship."
            ]
        },
        {
            'title': 'V. CONCLUSION',
            'paragraphs': [
                "For the foregoing reasons, Midwest Manufacturing Co. respectfully requests that this Court grant its Motion to Dismiss for lack of personal jurisdiction. In the alternative, Midwest requests that the Court transfer this action to the Southern District of Ohio pursuant to 28 U.S.C. § 1406(a).",
                "DATED: October 1, 2024",
                "Respectfully submitted,\n\nBAKER & HOSTETLER LLP\n\n/s/ Michael Roberts\nMichael Roberts (Ohio Bar No. 0078456)\n(Admitted Pro Hac Vice)\nCounsel for Defendant Midwest Manufacturing Co."
            ]
        }
    ]
}


if __name__ == "__main__":
    create_brief("01_mtd_contract_global_ventures.docx", brief1)
    create_brief("02_opp_summary_judgment_santos.docx", brief2)
    create_brief("03_preliminary_injunction_nexgen.docx", brief3)
    create_brief("04_mtd_personal_jurisdiction_midwest.docx", brief4)
    print("\nAll sample briefs generated successfully!")
