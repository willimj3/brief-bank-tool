"""
Document parsing for legal briefs.

Handles DOCX and PDF ingestion, extracting text while preserving
structure (headings, sections, footnotes) and identifying document
components (caption, facts, arguments, etc.).
"""

import re
import uuid
from pathlib import Path
from typing import Optional
from datetime import datetime

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
import pdfplumber

from .models import (
    Brief, BriefSection, ArgumentChunk, Citation,
    SectionType, ProceduralPosture
)


# Patterns for identifying brief sections
SECTION_PATTERNS = {
    SectionType.INTRODUCTION: [
        r"^I\.?\s*INTRODUCTION",
        r"^INTRODUCTION",
        r"^PRELIMINARY\s+STATEMENT",
    ],
    SectionType.STATEMENT_OF_FACTS: [
        r"^II\.?\s*STATEMENT\s+OF\s+FACTS",
        r"^STATEMENT\s+OF\s+FACTS",
        r"^FACTUAL\s+BACKGROUND",
        r"^FACTS",
        r"^BACKGROUND",
    ],
    SectionType.PROCEDURAL_HISTORY: [
        r"^PROCEDURAL\s+HISTORY",
        r"^PROCEDURAL\s+BACKGROUND",
    ],
    SectionType.LEGAL_STANDARD: [
        r"^LEGAL\s+STANDARD",
        r"^STANDARD\s+OF\s+REVIEW",
    ],
    SectionType.ARGUMENT: [
        r"^III\.?\s*ARGUMENT",
        r"^IV\.?\s*ARGUMENT",
        r"^ARGUMENT",
        r"^LEGAL\s+ARGUMENT",
        r"^DISCUSSION",
    ],
    SectionType.CONCLUSION: [
        r"^CONCLUSION",
        r"^V\.?\s*CONCLUSION",
        r"^VI\.?\s*CONCLUSION",
    ],
}

# Pattern for argument sub-sections (roman numerals, letters, numbers)
ARGUMENT_HEADING_PATTERNS = [
    r"^[IVX]+\.\s+.+",  # Roman numerals: I. , II. , etc.
    r"^[A-Z]\.\s+.+",   # Capital letters: A. , B. , etc.
    r"^\d+\.\s+.+",     # Numbers: 1. , 2. , etc.
]

# Citation pattern - matches standard legal citations
CITATION_PATTERN = re.compile(
    r"([A-Z][a-zA-Z\'\-\s]+(?:v\.|vs\.)\s+[A-Z][a-zA-Z\'\-\s]+),?\s*"
    r"(\d+)\s+"
    r"([A-Z][a-zA-Z\.\s\d]+?)\s+"
    r"(\d+)"
    r"(?:\s*,\s*(\d+))?"  # Optional pinpoint
    r"\s*\(([^)]+)\s+(\d{4})\)",
    re.MULTILINE
)

# Simpler citation pattern for cases we might miss
SIMPLE_CITATION_PATTERN = re.compile(
    r"([A-Z][a-zA-Z\'\-\s]+(?:v\.|vs\.)\s+[A-Z][a-zA-Z\'\-\s,]+\d{4}\))",
    re.MULTILINE
)


def extract_citations(text: str, chunk_id: str) -> list[Citation]:
    """Extract legal citations from text."""
    citations = []
    seen = set()

    for match in CITATION_PATTERN.finditer(text):
        full_text = match.group(0)
        if full_text in seen:
            continue
        seen.add(full_text)

        # Get context (surrounding sentence)
        start = max(0, match.start() - 100)
        end = min(len(text), match.end() + 100)
        context = text[start:end]

        citation = Citation(
            id=str(uuid.uuid4()),
            full_text=full_text,
            case_name=match.group(1).strip(),
            volume=match.group(2),
            reporter=match.group(3).strip(),
            page=match.group(4),
            pinpoint=match.group(5),
            court=match.group(6),
            year=int(match.group(7)) if match.group(7) else None,
            parent_chunk_id=chunk_id,
            context=context.strip()
        )
        citations.append(citation)

    return citations


def identify_section_type(heading: str) -> SectionType:
    """Identify what type of section a heading represents."""
    heading_upper = heading.upper().strip()

    for section_type, patterns in SECTION_PATTERNS.items():
        for pattern in patterns:
            if re.match(pattern, heading_upper):
                return section_type

    # Check if it's an argument sub-heading
    for pattern in ARGUMENT_HEADING_PATTERNS:
        if re.match(pattern, heading.strip()):
            return SectionType.ARGUMENT

    return SectionType.OTHER


def is_heading(paragraph: Paragraph) -> bool:
    """Check if a paragraph is likely a heading based on formatting."""
    if not paragraph.text.strip():
        return False

    # Check for explicit heading styles
    if paragraph.style and 'Heading' in paragraph.style.name:
        return True

    # Check for all caps (common in legal briefs)
    text = paragraph.text.strip()
    if text.isupper() and len(text) < 200:
        return True

    # Check for bold formatting
    if paragraph.runs:
        all_bold = all(run.bold for run in paragraph.runs if run.text.strip())
        if all_bold and len(text) < 200:
            return True

    # Check for roman numeral or letter patterns
    for pattern in ARGUMENT_HEADING_PATTERNS:
        if re.match(pattern, text):
            return True

    return False


def identify_procedural_posture(text: str) -> Optional[ProceduralPosture]:
    """Identify the procedural posture from brief text."""
    text_lower = text.lower()

    patterns = [
        (ProceduralPosture.MOTION_TO_DISMISS, [
            r"motion\s+to\s+dismiss",
            r"12\(b\)\(6\)",
            r"failure\s+to\s+state\s+a\s+claim"
        ]),
        (ProceduralPosture.SUMMARY_JUDGMENT, [
            r"motion\s+for\s+summary\s+judgment",
            r"summary\s+judgment",
            r"rule\s+56"
        ]),
        (ProceduralPosture.PRELIMINARY_INJUNCTION, [
            r"preliminary\s+injunction",
            r"temporary\s+restraining\s+order",
            r"tro"
        ]),
        (ProceduralPosture.MOTION_TO_COMPEL, [
            r"motion\s+to\s+compel"
        ]),
        (ProceduralPosture.OPPOSITION, [
            r"opposition\s+to",
            r"in\s+opposition"
        ]),
        (ProceduralPosture.REPLY, [
            r"reply\s+in\s+support",
            r"reply\s+brief",
            r"reply\s+memorandum"
        ]),
        (ProceduralPosture.APPEAL_BRIEF, [
            r"appellant[\'s]*\s+brief",
            r"appellee[\'s]*\s+brief",
            r"opening\s+brief",
            r"answering\s+brief"
        ]),
    ]

    for posture, posture_patterns in patterns:
        for pattern in posture_patterns:
            if re.search(pattern, text_lower):
                return posture

    return ProceduralPosture.OTHER


def extract_court_info(text: str) -> tuple[Optional[str], Optional[str]]:
    """Extract court and jurisdiction from brief text."""
    court = None
    jurisdiction = None

    # Federal district courts
    district_pattern = r"((?:Northern|Southern|Eastern|Western|Central|Middle)\s+District\s+of\s+\w+)"
    match = re.search(district_pattern, text, re.IGNORECASE)
    if match:
        court = match.group(1)
        jurisdiction = "federal"

    # Circuit courts
    circuit_pattern = r"(\d+(?:st|nd|rd|th)\s+Circuit|(?:First|Second|Third|Fourth|Fifth|Sixth|Seventh|Eighth|Ninth|Tenth|Eleventh|D\.?C\.?)\s+Circuit)"
    match = re.search(circuit_pattern, text, re.IGNORECASE)
    if match:
        court = match.group(1)
        jurisdiction = "federal"

    # State courts
    state_pattern = r"(?:Superior\s+Court|Supreme\s+Court|Court\s+of\s+Appeal)\s+of\s+(?:the\s+State\s+of\s+)?(\w+)"
    match = re.search(state_pattern, text, re.IGNORECASE)
    if match:
        court = match.group(0)
        jurisdiction = match.group(1).lower()

    return court, jurisdiction


def extract_case_info(text: str) -> tuple[Optional[str], Optional[str]]:
    """Extract case name and number from brief text."""
    case_name = None
    case_number = None

    # Case number patterns
    case_num_patterns = [
        r"Case\s+No\.?\s*:?\s*([\w\-:]+)",
        r"No\.?\s+([\d\-cv\w]+)",
        r"Docket\s+No\.?\s*:?\s*([\w\-]+)",
    ]

    for pattern in case_num_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            case_number = match.group(1)
            break

    # Case name - look for "v." pattern in caption area (first 2000 chars)
    caption_text = text[:2000]

    # Try multiple patterns for case name extraction
    case_name_patterns = [
        # Standard format: PLAINTIFF v. DEFENDANT (all caps)
        r"([A-Z][A-Z\s\.,\']+?)\s*,?\s*(?:Plaintiff|Petitioner|Appellant)s?\s*,?\s*v\.?\s+([A-Z][A-Z\s\.,\']+?)\s*,?\s*(?:Defendant|Respondent|Appellee)s?",
        # Mixed case with v.
        r"([A-Z][a-zA-Z\s\.,\'\-]+?)\s+v\.\s+([A-Z][a-zA-Z\s\.,\'\-]+?)(?:\s*,|\s*\n|\s*Case|\s*No\.)",
        # All caps with v.
        r"([A-Z][A-Z\s\.,\']+)\s+v\.\s+([A-Z][A-Z\s\.,\']+)",
    ]

    for pattern in case_name_patterns:
        match = re.search(pattern, caption_text, re.MULTILINE)
        if match:
            plaintiff = match.group(1).strip().rstrip(',').title()
            defendant = match.group(2).strip().rstrip(',').title()
            # Clean up extra whitespace
            plaintiff = ' '.join(plaintiff.split())
            defendant = ' '.join(defendant.split())
            # Truncate if too long (keep it readable)
            if len(plaintiff) > 40:
                plaintiff = plaintiff[:37] + "..."
            if len(defendant) > 40:
                defendant = defendant[:37] + "..."
            case_name = f"{plaintiff} v. {defendant}"
            break

    return case_name, case_number


def generate_brief_title(case_name: Optional[str], case_number: Optional[str], filename: str) -> str:
    """Generate a display title for a brief, preferring case name or number over filename."""
    if case_name:
        return case_name
    if case_number:
        return f"Case No. {case_number}"
    # Fall back to filename without extension and UUID-like prefixes
    name = Path(filename).stem
    # Remove UUID prefix if present (e.g., "abc123-def456-...-filename" -> "filename")
    if len(name) > 36 and name[8] == '-' and name[13] == '-':
        # Looks like it starts with a UUID, try to get the rest
        parts = name.split('-')
        if len(parts) > 5:
            name = '-'.join(parts[5:])
    return name if name else filename


def parse_docx(file_path: Path) -> Brief:
    """Parse a DOCX file into a Brief with sections and chunks."""
    doc = Document(file_path)
    brief_id = str(uuid.uuid4())

    # Collect all text for metadata extraction
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Extract metadata
    court, jurisdiction = extract_court_info(full_text)
    case_name, case_number = extract_case_info(full_text)
    procedural_posture = identify_procedural_posture(full_text)

    # Parse into sections
    sections = []
    current_section_type = SectionType.CAPTION
    current_section_text = []
    current_section_title = None
    section_order = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if is_heading(para):
            # Save previous section if exists
            if current_section_text:
                section = BriefSection(
                    id=str(uuid.uuid4()),
                    brief_id=brief_id,
                    section_type=current_section_type,
                    title=current_section_title,
                    content="\n".join(current_section_text),
                    order=section_order
                )
                sections.append(section)
                section_order += 1

            # Start new section
            current_section_type = identify_section_type(text)
            current_section_title = text
            current_section_text = []
        else:
            current_section_text.append(text)

    # Don't forget the last section
    if current_section_text:
        section = BriefSection(
            id=str(uuid.uuid4()),
            brief_id=brief_id,
            section_type=current_section_type,
            title=current_section_title,
            content="\n".join(current_section_text),
            order=section_order
        )
        sections.append(section)

    # Generate a good display title
    title = generate_brief_title(case_name, case_number, file_path.name)

    brief = Brief(
        id=brief_id,
        filename=file_path.name,
        title=title,
        court=court,
        jurisdiction=jurisdiction,
        case_name=case_name,
        case_number=case_number,
        procedural_posture=procedural_posture,
        sections=sections,
        full_text=full_text,
        file_type="docx"
    )

    return brief


def parse_pdf(file_path: Path) -> Brief:
    """Parse a PDF file into a Brief with sections and chunks."""
    brief_id = str(uuid.uuid4())

    full_text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                full_text += page_text + "\n"

    # Extract metadata
    court, jurisdiction = extract_court_info(full_text)
    case_name, case_number = extract_case_info(full_text)
    procedural_posture = identify_procedural_posture(full_text)

    # For PDFs, we do simpler section detection based on patterns
    sections = []
    lines = full_text.split("\n")
    current_section_type = SectionType.CAPTION
    current_section_text = []
    current_section_title = None
    section_order = 0

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Check if this line is a section heading
        detected_type = identify_section_type(line)
        if detected_type != SectionType.OTHER:
            # Save previous section
            if current_section_text:
                section = BriefSection(
                    id=str(uuid.uuid4()),
                    brief_id=brief_id,
                    section_type=current_section_type,
                    title=current_section_title,
                    content="\n".join(current_section_text),
                    order=section_order
                )
                sections.append(section)
                section_order += 1

            current_section_type = detected_type
            current_section_title = line
            current_section_text = []
        else:
            current_section_text.append(line)

    # Last section
    if current_section_text:
        section = BriefSection(
            id=str(uuid.uuid4()),
            brief_id=brief_id,
            section_type=current_section_type,
            title=current_section_title,
            content="\n".join(current_section_text),
            order=section_order
        )
        sections.append(section)

    # Generate a good display title
    title = generate_brief_title(case_name, case_number, file_path.name)

    brief = Brief(
        id=brief_id,
        filename=file_path.name,
        title=title,
        court=court,
        jurisdiction=jurisdiction,
        case_name=case_name,
        case_number=case_number,
        procedural_posture=procedural_posture,
        sections=sections,
        full_text=full_text,
        file_type="pdf"
    )

    return brief


def chunk_brief(brief: Brief) -> tuple[list[ArgumentChunk], list[Citation]]:
    """
    Break a brief into retrievable argument chunks.

    Strategy:
    - Each major argument section becomes a chunk
    - Sub-arguments are linked to parent chunks
    - Statement of facts is kept as a separate retrievable unit
    - Legal standards are tagged for high reusability
    """
    chunks = []
    citations = []

    for section in brief.sections:
        if section.section_type == SectionType.CAPTION:
            continue  # Skip caption, not useful for retrieval

        # For argument sections, try to break into sub-chunks by heading
        if section.section_type == SectionType.ARGUMENT:
            argument_chunks = _chunk_argument_section(
                section, brief
            )
            chunks.extend(argument_chunks)
        else:
            # Other sections become single chunks
            chunk = ArgumentChunk(
                id=str(uuid.uuid4()),
                brief_id=brief.id,
                section_type=section.section_type,
                heading=section.title,
                content=section.content,
                jurisdiction=brief.jurisdiction,
                court=brief.court,
                procedural_posture=brief.procedural_posture,
                is_legal_standard=(section.section_type == SectionType.LEGAL_STANDARD),
                is_factual=(section.section_type == SectionType.STATEMENT_OF_FACTS),
            )
            chunks.append(chunk)

        # Extract citations from all chunks
        for chunk in chunks:
            chunk_citations = extract_citations(chunk.content, chunk.id)
            chunk.citations = [c.id for c in chunk_citations]
            citations.extend(chunk_citations)

    return chunks, citations


def _chunk_argument_section(
    section: BriefSection,
    brief: Brief
) -> list[ArgumentChunk]:
    """Break an argument section into sub-chunks by heading."""
    chunks = []
    content = section.content
    lines = content.split("\n")

    current_heading = section.title
    current_text = []
    parent_chunk_id = None

    # Patterns for sub-headings within arguments
    subheading_patterns = [
        re.compile(r"^[A-Z]\.\s+.+"),  # A. , B. , etc.
        re.compile(r"^\d+\.\s+.+"),     # 1. , 2. , etc.
        re.compile(r"^[ivx]+\.\s+.+", re.IGNORECASE),  # roman numerals
    ]

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        is_subheading = any(p.match(line_stripped) for p in subheading_patterns)

        if is_subheading:
            # Save previous chunk
            if current_text:
                chunk = ArgumentChunk(
                    id=str(uuid.uuid4()),
                    brief_id=brief.id,
                    section_type=SectionType.ARGUMENT,
                    heading=current_heading,
                    content="\n".join(current_text),
                    parent_chunk_id=parent_chunk_id,
                    jurisdiction=brief.jurisdiction,
                    court=brief.court,
                    procedural_posture=brief.procedural_posture,
                )
                chunks.append(chunk)

                # First chunk becomes parent for subsequent chunks
                if parent_chunk_id is None:
                    parent_chunk_id = chunk.id

            current_heading = line_stripped
            current_text = []
        else:
            current_text.append(line_stripped)

    # Don't forget the last chunk
    if current_text:
        chunk = ArgumentChunk(
            id=str(uuid.uuid4()),
            brief_id=brief.id,
            section_type=SectionType.ARGUMENT,
            heading=current_heading,
            content="\n".join(current_text),
            parent_chunk_id=parent_chunk_id,
            jurisdiction=brief.jurisdiction,
            court=brief.court,
            procedural_posture=brief.procedural_posture,
        )
        chunks.append(chunk)

    return chunks


def parse_document(file_path: Path) -> Brief:
    """Parse a document (DOCX or PDF) into a Brief."""
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        return parse_docx(file_path)
    elif suffix == ".pdf":
        return parse_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
