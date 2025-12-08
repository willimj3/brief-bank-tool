"""
DOCX export functionality.

Generates properly formatted Word documents from draft briefs,
preserving legal document formatting conventions.
"""

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from .models import DraftBrief, GeneratedSection


def create_brief_document(draft: DraftBrief, output_path: Path) -> Path:
    """
    Export a draft brief to a properly formatted DOCX file.

    Formatting follows standard legal brief conventions:
    - 1-inch margins
    - 12pt Times New Roman
    - Double-spaced body text
    - Proper heading hierarchy
    """
    doc = Document()

    # Set up styles
    _setup_styles(doc)

    # Add caption/header
    _add_caption(doc, draft)

    # Add each generated section
    for section in draft.sections:
        _add_section(doc, section)

    # Add signature block placeholder
    _add_signature_block(doc)

    # Save document
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return output_path


def _setup_styles(doc: Document):
    """Set up document styles for legal brief formatting."""
    # Modify Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 2.0  # Double-spaced
    paragraph_format.space_after = Pt(0)

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Create Heading 1 style for main sections
    try:
        heading1 = doc.styles.add_style('BriefHeading1', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        heading1 = doc.styles['BriefHeading1']

    heading1.font.name = 'Times New Roman'
    heading1.font.size = Pt(12)
    heading1.font.bold = True
    heading1.font.all_caps = True
    heading1.paragraph_format.space_before = Pt(24)
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create Heading 2 style for sub-sections
    try:
        heading2 = doc.styles.add_style('BriefHeading2', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        heading2 = doc.styles['BriefHeading2']

    heading2.font.name = 'Times New Roman'
    heading2.font.size = Pt(12)
    heading2.font.bold = True
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)


def _add_caption(doc: Document, draft: DraftBrief):
    """Add the case caption to the document."""
    matter = draft.matter

    # Court name (centered, all caps)
    court_para = doc.add_paragraph()
    court_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = court_para.add_run(matter.court.upper())
    run.bold = True

    # Blank line
    doc.add_paragraph()

    # Case name and number in a simple format
    # Left side: plaintiff info
    case_para = doc.add_paragraph()
    case_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Simple case caption
    lines = [
        matter.case_name,
        "",
        f"Case No. [CASE NUMBER]",
        "",
        f"{matter.procedural_posture.value.upper().replace('_', ' ')}",
    ]

    for i, line in enumerate(lines):
        if i > 0:
            case_para.add_run("\n")
        run = case_para.add_run(line)
        if i == 0:  # Case name
            run.bold = True

    # Horizontal line
    doc.add_paragraph("_" * 60)
    doc.add_paragraph()


def _add_section(doc: Document, section: GeneratedSection):
    """Add a generated section to the document."""
    # Add heading
    heading_para = doc.add_paragraph(section.heading, style='BriefHeading1')

    # Add content
    # Split content by paragraphs and add each
    paragraphs = section.content.split('\n\n')
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        para = doc.add_paragraph(para_text)

        # Highlight [CITATION NEEDED] and [FACT PLACEHOLDER] markers
        # (In production, you'd use runs with highlighting)

    # Add warnings as comments/notes if present
    if section.warnings:
        warning_para = doc.add_paragraph()
        warning_para.add_run("[REVIEW NOTES: ").bold = True
        for warning in section.warnings:
            warning_para.add_run(f"• {warning} ")
        warning_para.add_run("]").bold = True

    doc.add_paragraph()  # Space after section


def _add_signature_block(doc: Document):
    """Add signature block placeholder."""
    doc.add_paragraph()
    doc.add_paragraph()

    sig_para = doc.add_paragraph()
    sig_para.add_run("Respectfully submitted,")

    doc.add_paragraph()
    doc.add_paragraph()

    sig_line = doc.add_paragraph()
    sig_line.add_run("_" * 40)

    attorney_para = doc.add_paragraph()
    attorney_para.add_run("[ATTORNEY NAME]")

    firm_para = doc.add_paragraph()
    firm_para.add_run("[FIRM NAME]")

    address_para = doc.add_paragraph()
    address_para.add_run("[ADDRESS]")

    date_para = doc.add_paragraph()
    date_para.add_run(f"Dated: {datetime.now().strftime('%B %d, %Y')}")


def export_draft(draft: DraftBrief, output_dir: Path = None) -> Path:
    """
    Export a draft brief to DOCX.

    Returns the path to the exported file.
    """
    if output_dir is None:
        output_dir = Path(__file__).parent.parent.parent / "data" / "exports"

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Generate filename from case name and timestamp
    case_name = draft.matter.case_name.replace(" ", "_").replace(".", "")[:30]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{case_name}_{timestamp}.docx"

    output_path = output_dir / filename

    return create_brief_document(draft, output_path)
